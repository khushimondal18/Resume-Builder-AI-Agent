from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os


def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_heading(doc, title):
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    add_bottom_border(p)


def add_bullet(doc, text):
    if text.strip():
        doc.add_paragraph(text.strip(), style="List Bullet")


def create_resume_docx(resume_data, final_resume_text=None):
    os.makedirs("outputs", exist_ok=True)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    # HEADER: left/right layout
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    p = left.paragraphs[0]
    name = p.add_run(resume_data.get("name", "").upper())
    name.bold = True
    name.font.size = Pt(16)

    if resume_data.get("roll_no"):
        left.add_paragraph(f"Roll No.: {resume_data.get('roll_no')}")
    if resume_data.get("degree"):
        left.add_paragraph(resume_data.get("degree"))
    if resume_data.get("institute"):
        left.add_paragraph(resume_data.get("institute"))

    contact_lines = [
        resume_data.get("phone", ""),
        resume_data.get("email", ""),
        resume_data.get("github", ""),
        resume_data.get("linkedin", ""),
    ]

    for item in contact_lines:
        if item.strip():
            p = right.add_paragraph(item)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # EDUCATION TABLE
    education = resume_data.get("education_data", [])
    if education:
        add_heading(doc, "Education")

        edu_table = doc.add_table(rows=1, cols=4)
        edu_table.style = "Table Grid"

        headers = ["Degree / Certificate", "Institute / School", "CGPA / Percentage", "Year"]
        for i, h in enumerate(headers):
            run = edu_table.rows[0].cells[i].paragraphs[0].add_run(h)
            run.bold = True

        for edu in education:
            row = edu_table.add_row().cells
            row[0].text = edu.get("degree", "")
            row[1].text = edu.get("institute", "")
            row[2].text = edu.get("score", "")
            row[3].text = edu.get("year", "")

    # EXPERIENCE
    experience = resume_data.get("experience_data", [])
    if experience:
        add_heading(doc, "Experience")

        for exp in experience:
            p = doc.add_paragraph()
            run = p.add_run(f"• {exp.get('company', '')}")
            run.bold = True

            if exp.get("duration"):
                p.add_run(f"                                      {exp.get('duration')}")

            if exp.get("role") or exp.get("location"):
                p = doc.add_paragraph()
                p.add_run(exp.get("role", "")).italic = True
                if exp.get("location"):
                    p.add_run(f" | {exp.get('location')}")

            for bullet in exp.get("bullets", []):
                add_bullet(doc, bullet)

    # PROJECTS
    projects = resume_data.get("project_data", [])
    if projects:
        add_heading(doc, "Projects")

        for project in projects:
            p = doc.add_paragraph()
            run = p.add_run(f"• {project.get('name', '')}")
            run.bold = True

            if project.get("tech_stack"):
                p.add_run(f" | {project.get('tech_stack')}")

            if project.get("duration"):
                p.add_run(f"                                      {project.get('duration')}")

            if project.get("link"):
                doc.add_paragraph(project.get("link"))

            for bullet in project.get("bullets", []):
                add_bullet(doc, bullet)

    # TECHNICAL SKILLS
    skills_items = {
        "Programming Languages": resume_data.get("programming", ""),
        "Frameworks / Libraries": resume_data.get("frameworks", ""),
        "Databases": resume_data.get("databases", ""),
        "Tools": resume_data.get("tools", ""),
        "Cloud / Platforms": resume_data.get("cloud", ""),
        "Soft Skills": resume_data.get("soft_skills", ""),
    }

    non_empty_skills = {k: v for k, v in skills_items.items() if v.strip()}

    if non_empty_skills:
        add_heading(doc, "Technical Skills")

        for key, value in non_empty_skills.items():
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(value)

    # COURSES
    if resume_data.get("courses", "").strip():
        add_heading(doc, "Key Courses Taken")
        for course in resume_data.get("courses", "").split("\n"):
            add_bullet(doc, course)

    # ACHIEVEMENTS
    if resume_data.get("achievements", "").strip():
        add_heading(doc, "Achievements")
        for achievement in resume_data.get("achievements", "").split("\n"):
            add_bullet(doc, achievement.strip("-• "))

    # POSITIONS
    if resume_data.get("positions", "").strip():
        add_heading(doc, "Positions of Responsibility")
        for position in resume_data.get("positions", "").split("\n"):
            add_bullet(doc, position.strip("-• "))

    # EXTRACURRICULAR
    if resume_data.get("extra", "").strip():
        add_heading(doc, "Extracurricular Activities")
        for item in resume_data.get("extra", "").split("\n"):
            add_bullet(doc, item.strip("-• "))
    custom_sections = resume_data.get("custom_sections", {})
    for section_name, items in custom_sections.items():
        if items:
            add_heading(doc, section_name)
            for item in items:
                add_bullet(doc, item)
    file_path = "outputs/generated_resume.docx"
    doc.save(file_path)
    return file_path